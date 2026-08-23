from io import BytesIO
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------- helpers
def _set_font(run, size=11, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    # ép font cho cả phần chữ có dấu tiếng Việt (đông á)
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "Times New Roman")


def _p(doc, text="", size=11, bold=False, italic=False, align=None):
    para = doc.add_paragraph()
    if align:
        para.alignment = align
    run = para.add_run(text)
    _set_font(run, size, bold, italic)
    return para


def _cell_text(cell, text, size=10.5, bold=False):
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(str(text) if text not in (None, "") else "—")
    _set_font(run, size, bold)


def _field_table(doc, rows):
    """rows: list các tuple (nhãn, giá trị) -> bảng 2 cột, không viền, giống mẫu giấy."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = True
    table.columns[0].width = Cm(5.5)
    for i, (label, value) in enumerate(rows):
        _cell_text(table.cell(i, 0), label, bold=True)
        _cell_text(table.cell(i, 1), value)
    return table


def _section_heading(doc, number, title):
    _p(doc, f"{number}. {title}", size=12, bold=True)


def _list_table(doc, headers, rows):
    """Bảng có viền cho các mục lặp (đào tạo, khen thưởng...)."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        _cell_text(table.rows[0].cells[i], h, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            _cell_text(cells[i], val)
    doc.add_paragraph()
    return table


def _ngay(d):
    return d.strftime("%d/%m/%Y") if d else ""


# ---------------------------------------------------------------- hàm chính
def tao_docx_so_yeu_ly_lich(cb):
    doc = Document()

    # lề trang theo chuẩn văn bản hành chính
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    # --- Quốc hiệu tiêu ngữ ---
    _p(doc, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(doc, "Độc lập - Tự do - Hạnh phúc", size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(doc, "—————————", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    _p(doc, "SƠ YẾU LÝ LỊCH CÁN BỘ, CÔNG CHỨC", size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(doc, "(Mẫu 2C-TCTW/98)", size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    # --- 1. Thông tin cá nhân ---
    _section_heading(doc, 1, "Thông tin cá nhân")
    _field_table(doc, [
        ("Họ và tên khai sinh", cb.ho_ten_khai_sinh),
        ("Tên gọi khác", cb.ten_goi_khac),
        ("Giới tính", cb.gioi_tinh),
        ("Ngày sinh", _ngay(cb.ngay_sinh)),
        ("Nơi sinh", cb.noi_sinh),
        ("Quê quán", ", ".join(filter(None, [cb.que_quan_xa, cb.que_quan_huyen, cb.que_quan_tinh]))),
        ("Nơi ở hiện nay", cb.noi_o_hien_nay),
        ("Điện thoại", cb.dien_thoai),
        ("Số CMND/CCCD", cb.so_cmnd_cccd),
        ("Dân tộc", cb.dan_toc),
        ("Tôn giáo", cb.ton_giao),
        ("Thành phần gia đình xuất thân", cb.thanh_phan_gia_dinh),
    ])
    doc.add_paragraph()

    # --- 2. Tuyển dụng, Đảng, đoàn thể, quân đội ---
    _section_heading(doc, 2, "Tuyển dụng, Đảng, đoàn thể, quân đội")
    _field_table(doc, [
        ("Ngày tuyển dụng", _ngay(cb.ngay_tuyen_dung)),
        ("Nơi tuyển dụng", cb.noi_tuyen_dung),
        ("Ngày vào cơ quan hiện tại", _ngay(cb.ngay_vao_co_quan_hien_tai)),
        ("Ngày vào Đảng", _ngay(cb.ngay_vao_dang)),
        ("Ngày chính thức", _ngay(cb.ngay_chinh_thuc_dang)),
        ("Ngày nhập ngũ", _ngay(cb.ngay_nhap_ngu)),
        ("Ngày xuất ngũ", _ngay(cb.ngay_xuat_ngu)),
        ("Quân hàm / chức vụ cao nhất", cb.quan_ham_chuc_vu_cao_nhat),
    ])
    doc.add_paragraph()

    # --- 3. Trình độ, chức vụ, ngạch lương ---
    _section_heading(doc, 3, "Trình độ, chức vụ, ngạch lương")
    _field_table(doc, [
        ("Học hàm / học vị cao nhất", cb.hoc_ham_hoc_vi_cao_nhat),
        ("Lý luận chính trị", cb.ly_luan_chinh_tri),
        ("Ngoại ngữ", cb.ngoai_ngu),
        ("Chức vụ", cb.chuc_vu_ref.ten_chuc_vu if cb.chuc_vu_ref else ""),
        ("Đơn vị công tác", cb.don_vi.ten_don_vi if cb.don_vi else ""),
        ("Ngạch công chức", cb.ngach_cong_chuc),
        ("Mã số ngạch", cb.ma_so_ngach),
        ("Bậc lương / Hệ số", f"{cb.bac_luong or ''} / {cb.he_so_luong or ''}"),
    ])
    doc.add_paragraph()

    # --- 26. Đào tạo, bồi dưỡng ---
    _section_heading(doc, 26, "Đào tạo, bồi dưỡng")
    if cb.dao_tao:
        _list_table(
            doc,
            ["Tên trường", "Ngành học", "Từ năm", "Đến năm", "Hình thức", "Văn bằng"],
            [[d.ten_truong, d.nganh_hoc, d.tu_nam, d.den_nam, d.hinh_thuc_hoc, d.van_bang_chung_chi] for d in cb.dao_tao],
        )
    else:
        _p(doc, "Không có dữ liệu.", italic=True)
    doc.add_paragraph()

    # --- 27. Quá trình công tác ---
    _section_heading(doc, 27, "Tóm tắt quá trình công tác")
    if cb.qua_trinh_cong_tac:
        _list_table(
            doc,
            ["Từ", "Đến", "Chức danh, chức vụ, đơn vị công tác"],
            [[q.tu_thang_nam, q.den_thang_nam, q.chuc_danh_don_vi] for q in cb.qua_trinh_cong_tac],
        )
    else:
        _p(doc, "Không có dữ liệu.", italic=True)
    doc.add_paragraph()

    # --- 22-23. Khen thưởng / Kỷ luật ---
    _section_heading(doc, 22, "Khen thưởng")
    if cb.khen_thuong:
        for k in cb.khen_thuong:
            _p(doc, f"• {k.noi_dung} ({k.nam})")
    else:
        _p(doc, "Không có.", italic=True)

    _section_heading(doc, 23, "Kỷ luật")
    if cb.ky_luat:
        for k in cb.ky_luat:
            _p(doc, f"• {k.hinh_thuc} — {k.ly_do} ({k.nam}, {k.cap_quyet_dinh})")
    else:
        _p(doc, "Không có.", italic=True)
    doc.add_paragraph()

    # --- 30. Quan hệ gia đình ---
    _section_heading(doc, 30, "Quan hệ gia đình")
    if cb.quan_he_gia_dinh:
        _list_table(
            doc,
            ["Nhóm", "Quan hệ", "Họ tên", "Năm sinh", "Thông tin"],
            [[g.nhom, g.quan_he, g.ho_ten, g.nam_sinh, g.thong_tin] for g in cb.quan_he_gia_dinh],
        )
    else:
        _p(doc, "Không có dữ liệu.", italic=True)
    doc.add_paragraph()

    # --- Chữ ký ---
    doc.add_paragraph()
    _p(doc, "..........., ngày ..... tháng ..... năm ..........", italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _p(doc, "Người khai ký tên", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer